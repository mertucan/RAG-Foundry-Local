from __future__ import annotations

from pathlib import Path


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx"}


def load_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8")
    if suffix == ".docx":
        return _load_docx(path)
    raise ValueError(f"Desteklenmeyen dosya turu: {path.suffix}")


def _load_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)
